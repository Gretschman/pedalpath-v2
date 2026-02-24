"""
export_doc.py
Converts a markdown file to a formatted .docx with Rob's exact header/footer spec.

Spec:
    Header left:  filename, 10pt Arial bold
    Header right: date/time mm/dd/yy hh:mm, 10pt Arial
    Footer left:  full file path, 8pt Arial, #888888
    Footer right: p. N of NN, 10pt Arial
    Body font:    11pt Arial
    Page:         US Letter, 1-inch margins
    H1:           16pt Arial bold, #1F3864
    H2:           13pt Arial bold, #2E75B6

Usage:
    python3 tools/export_doc.py path/to/file.md
    python3 tools/export_doc.py path/to/file.md --out path/to/output.docx
    python3 tools/export_doc.py docs/generated/session_log.md
"""

import sys
import argparse
import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants
except ImportError:
    print('[FAIL]  python-docx not installed. Run: pip install python-docx')
    sys.exit(1)

# ── Colors ────────────────────────────────────────────────────────────────────
COLOR_H1      = RGBColor(0x1F, 0x38, 0x64)
COLOR_H2      = RGBColor(0x2E, 0x75, 0xB6)
COLOR_FOOTER  = RGBColor(0x88, 0x88, 0x88)

FONT = 'Arial'


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_run_font(run, size_pt: float, bold=False, color: RGBColor | None = None):
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_paragraph(doc: Document, text: str, style=None) -> object:
    p = doc.add_paragraph(style=style)
    return p, p.add_run(text)


def set_page_size(section):
    """US Letter, 1-inch margins."""
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)


def make_field_element(field_type: str) -> OxmlElement:
    """Create a Word fldChar/instrText combo for PAGE or NUMPAGES fields."""
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' {field_type} '

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    return fld, instr, fld_end


def add_page_number_field(run, field: str):
    """Inject a PAGE or NUMPAGES field into a run."""
    r_elem = run._r
    fld, instr, fld_end = make_field_element(field)
    r_new_begin = OxmlElement('w:r')
    r_new_begin.append(fld)
    r_new_instr = OxmlElement('w:r')
    r_new_instr.append(instr)
    r_new_end = OxmlElement('w:r')
    r_new_end.append(fld_end)

    p = r_elem.getparent()
    idx = list(p).index(r_elem)
    p.insert(idx + 1, r_new_end)
    p.insert(idx + 1, r_new_instr)
    p.insert(idx + 1, r_new_begin)


def build_header(section, filename: str, timestamp: str):
    """Header: filename (left) | date (right) on same line using a tab stop."""
    header = section.header
    header.is_linked_to_previous = False

    # Clear default paragraph
    for p in header.paragraphs:
        p.clear()
    hdr_para = header.paragraphs[0]

    # Tab stop at right margin (6.5" text width → right-align tab at 6.5")
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = hdr_para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9360')   # 6.5" × 1440 twips/inch = 9360
    tabs.append(tab)
    pPr.append(tabs)

    # Left: filename bold
    run_left = hdr_para.add_run(filename)
    set_run_font(run_left, 10, bold=True)

    # Tab
    hdr_para.add_run('\t')

    # Right: timestamp
    run_right = hdr_para.add_run(timestamp)
    set_run_font(run_right, 10)


def build_footer(section, filepath: str):
    """Footer: full path (left, grey 8pt) | p. N of NN (right)."""
    footer = section.footer
    footer.is_linked_to_previous = False

    for p in footer.paragraphs:
        p.clear()
    ftr_para = footer.paragraphs[0]

    # Right-align tab stop at 6.5"
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = ftr_para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9360')
    tabs.append(tab)
    pPr.append(tabs)

    # Left: file path, grey 8pt
    run_path = ftr_para.add_run(filepath)
    set_run_font(run_path, 8, color=COLOR_FOOTER)

    # Tab
    ftr_para.add_run('\t')

    # Right: "p. " + PAGE + " of " + NUMPAGES
    run_p = ftr_para.add_run('p. ')
    set_run_font(run_p, 10)

    # PAGE field
    run_page = ftr_para.add_run()
    set_run_font(run_page, 10)
    add_page_number_field(run_page, 'PAGE')

    run_of = ftr_para.add_run(' of ')
    set_run_font(run_of, 10)

    # NUMPAGES field
    run_total = ftr_para.add_run()
    set_run_font(run_total, 10)
    add_page_number_field(run_total, 'NUMPAGES')


# ── Markdown parser ────────────────────────────────────────────────────────────

def parse_markdown_to_doc(doc: Document, md_text: str):
    """
    Very lightweight markdown → docx converter.
    Handles: # H1, ## H2, ### H3, ``` code blocks, - bullet, blank lines.
    Inline: **bold**, `code`, plain text.
    """
    in_code_block = False
    code_lines = []

    def flush_code(lines):
        if not lines:
            return
        p = doc.add_paragraph()
        run = p.add_run('\n'.join(lines))
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.25)

    def add_inline(para, text: str):
        """Parse inline **bold** and `code` markers."""
        pattern = r'(\*\*.*?\*\*|`.*?`)'
        parts = re.split(pattern, text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.font.name = FONT
                run.font.size = Pt(11)
                run.font.bold = True
            elif part.startswith('`') and part.endswith('`'):
                run = para.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            else:
                run = para.add_run(part)
                run.font.name = FONT
                run.font.size = Pt(11)

    for raw_line in md_text.splitlines():
        line = raw_line.rstrip()

        # Code block toggle
        if line.startswith('```'):
            if in_code_block:
                flush_code(code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Headings
        if line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            set_run_font(run, 12, bold=True, color=COLOR_H2)
            p.paragraph_format.space_before = Pt(6)
            continue

        if line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            set_run_font(run, 13, bold=True, color=COLOR_H2)
            p.paragraph_format.space_before = Pt(10)
            continue

        if line.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            set_run_font(run, 16, bold=True, color=COLOR_H1)
            p.paragraph_format.space_before = Pt(14)
            continue

        # Horizontal rule
        if re.match(r'^-{3,}$', line) or re.match(r'^_{3,}$', line):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            continue

        # Bullet list
        if line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, line[2:])
            for run in p.runs:
                if not run.font.size:
                    run.font.size = Pt(11)
            continue

        # Blank line
        if not line.strip():
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_inline(p, line)

    # Flush any unclosed code block
    if in_code_block and code_lines:
        flush_code(code_lines)


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Export markdown → formatted docx.')
    parser.add_argument('input', help='Input markdown file path')
    parser.add_argument('--out', help='Output .docx path (default: same name as input)')
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    if not input_path.exists():
        print(f'[FAIL]  File not found: {input_path}')
        sys.exit(1)

    if args.out:
        output_path = Path(args.out).resolve()
    else:
        output_path = input_path.with_suffix('.docx')

    print()
    print('=' * 60)
    print('  PedalPath v2 — Export Doc')
    print('=' * 60)
    print()
    print(f'  Input:  {input_path}')
    print(f'  Output: {output_path}')
    print()

    md_text   = input_path.read_text(encoding='utf-8')
    filename  = input_path.name
    filepath  = str(input_path)
    timestamp = datetime.now().strftime('%m/%d/%y %H:%M')

    doc = Document()
    section = doc.sections[0]
    set_page_size(section)

    # Set default body font
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(11)

    # Header + footer
    build_header(section, filename, timestamp)
    build_footer(section, filepath)

    # Body content
    parse_markdown_to_doc(doc, md_text)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    print(f'[PASS]  Written: {output_path}')
    print()


if __name__ == '__main__':
    main()
