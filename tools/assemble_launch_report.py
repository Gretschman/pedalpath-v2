"""
assemble_launch_report.py
Assembles the PedalPath Launch Roadmap Report as a formatted DOCX.

Sources:
    - docs/generated/launch_roadmap_report.md  (content)
    - _OUTPUT/PedalPath_PhaseRoadmap_*.png     (Diagram A)
    - _OUTPUT/PedalPath_Architecture_*.png     (Diagram B)

Output:
    /mnt/c/Users/Rob/Dropbox/!PedalPath/_OUTPUT/
    PedalPath_LaunchRoadmap_CriticalIssues_2026-03-24.docx

Usage:
    python3 tools/assemble_launch_report.py
    python3 tools/assemble_launch_report.py --regen-diagrams
"""

import sys
import re
import argparse
import subprocess
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print('[FAIL]  python-docx not installed. Run: pip install python-docx')
    sys.exit(1)

# ── Paths ──────────────────────────────────────────────────────────────────────
REPO_ROOT   = Path(__file__).parent.parent
MD_SOURCE   = REPO_ROOT / 'docs' / 'generated' / 'launch_roadmap_report.md'
OUTPUT_DIR  = Path('/mnt/c/Users/Rob/Dropbox/!PedalPath/_OUTPUT')
DATE_STR    = '2026-03-24'
DOCX_OUT    = OUTPUT_DIR / f'PedalPath_LaunchRoadmap_CriticalIssues_{DATE_STR}.docx'
DIAGRAM_A   = OUTPUT_DIR / f'PedalPath_PhaseRoadmap_{DATE_STR}.png'
DIAGRAM_B   = OUTPUT_DIR / f'PedalPath_Architecture_BeforeAfter_{DATE_STR}.png'
DIAG_SCRIPT = REPO_ROOT / 'tools' / 'generate_report_diagrams.py'

# ── Colors ─────────────────────────────────────────────────────────────────────
COLOR_H1     = RGBColor(0x1F, 0x38, 0x64)   # navy
COLOR_H2     = RGBColor(0x2E, 0x75, 0xB6)   # blue
COLOR_H3     = RGBColor(0x1F, 0x78, 0x4E)   # green
COLOR_FOOTER = RGBColor(0x88, 0x88, 0x88)
COLOR_COVER  = RGBColor(0x1F, 0x38, 0x64)
FONT = 'Arial'


# ── Font helpers ───────────────────────────────────────────────────────────────

def set_run_font(run, size_pt, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_inline(para, text: str, base_size=11):
    """Parse inline **bold**, `code`, and plain text."""
    pattern = r'(\*\*.*?\*\*|`.*?`)'
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_run_font(run, base_size, bold=True)
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            run = para.add_run(part)
            set_run_font(run, base_size)


# ── Page setup helpers ─────────────────────────────────────────────────────────

def set_page_size(section):
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)


def make_field_element(field_type: str):
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' {field_type} '
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    return fld, instr, fld_end


def add_page_number_field(run, field: str):
    r_elem = run._r
    fld, instr, fld_end = make_field_element(field)
    r_new_begin = OxmlElement('w:r'); r_new_begin.append(fld)
    r_new_instr = OxmlElement('w:r'); r_new_instr.append(instr)
    r_new_end   = OxmlElement('w:r'); r_new_end.append(fld_end)
    p = r_elem.getparent()
    idx = list(p).index(r_elem)
    p.insert(idx + 1, r_new_end)
    p.insert(idx + 1, r_new_instr)
    p.insert(idx + 1, r_new_begin)


def build_header(section, filename: str, timestamp: str):
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()
    hdr_para = header.paragraphs[0]
    pPr = hdr_para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9360')
    tabs.append(tab)
    pPr.append(tabs)
    run_left = hdr_para.add_run(filename)
    set_run_font(run_left, 10, bold=True)
    hdr_para.add_run('\t')
    run_right = hdr_para.add_run(timestamp)
    set_run_font(run_right, 10)


def build_footer(section, filepath: str):
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    ftr_para = footer.paragraphs[0]
    pPr = ftr_para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9360')
    tabs.append(tab)
    pPr.append(tabs)
    run_path = ftr_para.add_run(filepath)
    set_run_font(run_path, 8, color=COLOR_FOOTER)
    ftr_para.add_run('\t')
    run_p = ftr_para.add_run('p. ')
    set_run_font(run_p, 10)
    run_page = ftr_para.add_run()
    set_run_font(run_page, 10)
    add_page_number_field(run_page, 'PAGE')
    run_of = ftr_para.add_run(' of ')
    set_run_font(run_of, 10)
    run_total = ftr_para.add_run()
    set_run_font(run_total, 10)
    add_page_number_field(run_total, 'NUMPAGES')


def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_break_type())


def docx_break_type():
    from docx.enum.text import WD_BREAK
    return WD_BREAK.PAGE


# ── Cover page ─────────────────────────────────────────────────────────────────

def add_cover_page(doc):
    # Big vertical space
    for _ in range(6):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PedalPath v2')
    set_run_font(run, 32, bold=True, color=COLOR_H1)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Launch Roadmap & Critical Issues Report')
    set_run_font(run2, 22, bold=True, color=COLOR_H1)

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('Exhaustion-Mode Reference Guide — Print Before Acting')
    set_run_font(run3, 14, italic=True, color=COLOR_H2)

    doc.add_paragraph()
    doc.add_paragraph()

    for line, size in [
        ('Date: March 24, 2026', 12),
        ('Version: v1.0', 12),
        ('Completion estimate: ~12% to first revenue', 12),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_run_font(run, size, color=RGBColor(80, 80, 80))

    doc.add_paragraph()
    doc.add_paragraph()

    p_box = doc.add_paragraph()
    p_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_box = p_box.add_run('If you only read one page: go to PART 8 (Rob\'s Action List)')
    set_run_font(run_box, 13, bold=True, color=RGBColor(180, 60, 0))

    # Page break
    doc.add_page_break()


# ── Checkbox paragraph ─────────────────────────────────────────────────────────

def add_checkbox(doc, text: str):
    """Add a checkbox-style list item: '☐ text'."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run_box = p.add_run('\u25a1 ')  # ☐ unicode square
    set_run_font(run_box, 11, bold=True, color=COLOR_H2)
    add_inline(p, text)


# ── Table renderer ─────────────────────────────────────────────────────────────

def add_markdown_table(doc, lines_block: list[str]):
    """Render a simple markdown table into a docx table."""
    rows = []
    for line in lines_block:
        if re.match(r'^\s*\|[-| ]+\|\s*$', line):
            continue  # separator row
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        if cells:
            rows.append(cells)
    if not rows:
        return

    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        row = tbl.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= ncols:
                break
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            is_header = (i == 0)
            set_run_font(run, 10, bold=is_header,
                         color=COLOR_H1 if is_header else None)

    doc.add_paragraph()  # spacer after table


# ── Main markdown parser ────────────────────────────────────────────────────────

DIAGRAM_A_MARKER = '[DIAGRAM A — Phase Roadmap — inserted here in DOCX]'
DIAGRAM_B_MARKER = '[DIAGRAM B — Architecture Before/After — inserted here in DOCX]'


def parse_markdown_to_doc(doc, md_text: str, diagram_a: Path, diagram_b: Path):
    in_code  = False
    code_buf = []
    in_table = False
    table_buf= []

    def flush_code(lines):
        if not lines:
            return
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run('\n'.join(lines))
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    def flush_table(lines):
        add_markdown_table(doc, lines)

    raw_lines = md_text.splitlines()
    i = 0
    while i < len(raw_lines):
        line = raw_lines[i].rstrip()

        # Code block
        if line.startswith('```'):
            if in_code:
                flush_code(code_buf)
                code_buf = []
                in_code = False
            else:
                if in_table:
                    flush_table(table_buf)
                    table_buf = []
                    in_table = False
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Table detection
        if '|' in line and line.strip().startswith('|'):
            in_table = True
            table_buf.append(line)
            i += 1
            continue
        else:
            if in_table:
                flush_table(table_buf)
                table_buf = []
                in_table = False

        # Blank
        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^-{3,}$', line) or re.match(r'^_{3,}$', line):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Diagram markers
        stripped = line.strip().lstrip('*').strip()
        if DIAGRAM_A_MARKER in line or DIAGRAM_A_MARKER in stripped:
            if diagram_a and diagram_a.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = doc.add_picture(str(diagram_a), width=Inches(6.5))
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_cap = cap.add_run('Diagram A — Phase Roadmap')
                set_run_font(run_cap, 10, italic=True, color=COLOR_FOOTER)
            else:
                p = doc.add_paragraph('[Diagram A: Phase Roadmap — file not found]')
            i += 1
            continue

        if DIAGRAM_B_MARKER in line or DIAGRAM_B_MARKER in stripped:
            if diagram_b and diagram_b.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_picture(str(diagram_b), width=Inches(6.5))
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_cap = cap.add_run('Diagram B — Architecture Before / After')
                set_run_font(run_cap, 10, italic=True, color=COLOR_FOOTER)
            else:
                p = doc.add_paragraph('[Diagram B: Architecture Before/After — file not found]')
            i += 1
            continue

        # H1
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            set_run_font(run, 16, bold=True, color=COLOR_H1)
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after  = Pt(4)
            # Page break before major PART headers (but not the very first)
            if line.startswith('# PART') or line.startswith('# PedalPath'):
                pass  # cover page already handles first break
            i += 1
            continue

        # H2
        if line.startswith('## ') and not line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            set_run_font(run, 13, bold=True, color=COLOR_H2)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(3)
            i += 1
            continue

        # H3
        if line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            set_run_font(run, 12, bold=True, color=COLOR_H3)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            i += 1
            continue

        # Checkbox (lines starting with - [ ])
        if re.match(r'^- \[ \]', line):
            text = line[6:].strip()
            add_checkbox(doc, text)
            i += 1
            continue

        # Bullet list (- or *)
        if re.match(r'^[-*] ', line) or re.match(r'^  [-*] ', line):
            indent = 0.25 if line.startswith('  ') else 0
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25 + indent)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run_bullet = p.add_run('• ')
            set_run_font(run_bullet, 11, bold=False, color=COLOR_H2)
            text = re.sub(r'^  [-*] |^[-*] ', '', line)
            add_inline(p, text)
            i += 1
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m:
            num, text = m.group(1), m.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run_num = p.add_run(f'{num}. ')
            set_run_font(run_num, 11, bold=True, color=COLOR_H2)
            add_inline(p, text)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        add_inline(p, line)
        i += 1

    # Flush any open buffers
    if in_code and code_buf:
        flush_code(code_buf)
    if in_table and table_buf:
        flush_table(table_buf)


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Assemble Launch Roadmap DOCX')
    parser.add_argument('--regen-diagrams', action='store_true',
                        help='Regenerate PNG diagrams even if they exist')
    args = parser.parse_args()

    print()
    print('=' * 60)
    print('  PedalPath v2 — Assemble Launch Report DOCX')
    print('=' * 60)
    print()

    # 1. Ensure diagrams exist
    need_diagrams = args.regen_diagrams or not DIAGRAM_A.exists() or not DIAGRAM_B.exists()
    if need_diagrams:
        print('[RUN]  Generating diagrams...')
        result = subprocess.run([sys.executable, str(DIAG_SCRIPT)], capture_output=True, text=True)
        print(result.stdout)
        if result.returncode != 0:
            print(f'[WARN] Diagram generation had errors:\n{result.stderr}')
    else:
        print('[OK]   Diagrams already exist, skipping generation.')

    # 2. Read markdown
    if not MD_SOURCE.exists():
        print(f'[FAIL] Markdown source not found: {MD_SOURCE}')
        sys.exit(1)

    md_text   = MD_SOURCE.read_text(encoding='utf-8')
    filename  = DOCX_OUT.name
    filepath  = str(DOCX_OUT)
    timestamp = datetime.now().strftime('%m/%d/%y %H:%M')

    # 3. Build DOCX
    doc = Document()
    section = doc.sections[0]
    set_page_size(section)

    # Default body style
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(11)

    # Header + footer
    build_header(section, filename, timestamp)
    build_footer(section, filepath)

    # Cover page
    add_cover_page(doc)

    # Body
    parse_markdown_to_doc(doc, md_text, DIAGRAM_A, DIAGRAM_B)

    # 4. Save
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_OUT))

    print()
    print(f'[PASS]  DOCX written: {DOCX_OUT}')
    print()
    print('Verification checklist:')
    print('  [ ] Open in Word → confirm headers/footers visible')
    print('  [ ] Confirm both diagrams embedded (Diagram A, Diagram B)')
    print('  [ ] Print preview → readable at 100%, diagrams not clipped')
    print('  [ ] All 10 PARTs present in document')
    print()


if __name__ == '__main__':
    main()
